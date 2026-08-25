"""Real-PostgreSQL vertical evidence for the Personal Knowledge import loop.

RC-01 §D: actual DocumentConversionService behavior (no fake converter) for
PDF / DOCX / Markdown fixtures with fixed Chinese and English markers; the
durable two-phase worker; canonical Markdown, heading_path, position, segment
evidence, table preservation; browser search citations; archived documents
disappearing from search.
"""

from __future__ import annotations

import uuid

import pytest
from sqlalchemy import select

from app.database import Base
from app.models.knowledge import KnowledgeDocument, KnowledgeSegment
from app.models.tenant import Tenant
from app.models.user import User
from app.services.personal_knowledge_access import HumanBrowserPrincipal
from app.services.personal_knowledge_service import PersonalKnowledgeService

MARKER_EN = "WEEKEND-RC-20260825-PKB-EN-MARKER"
MARKER_ZH = "周末RC20260825个人知识唯一标记"
TABLE_MARKER = "TABLE-CELL-MARKER-8848"
PDF_MARKER = "WEEKEND-RC-20260825-PDF-ONLY-MARKER"
DOCX_MARKER = "WEEKEND-RC-20260825-DOCX-ONLY-MARKER"
TXT_MARKER = "WEEKEND-RC-20260825-TXT-ONLY-MARKER"
TXT_CELL = "TXT-CELL-MARKER-5516"


def _minimal_pdf_bytes(lines: list[str]) -> bytes:
    """Hand-crafted minimal PDF (standard Helvetica; ASCII text only)."""
    content = [f"BT /F1 12 Tf 72 {720 - 20 * index} Td ({line}) Tj ET" for index, line in enumerate(lines)]
    objs = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
    ]
    stream = "\n".join(content)
    objs.append(f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream")
    objs.append("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    pdf = b"%PDF-1.4\n"
    offsets: list[int] = []
    for index, obj in enumerate(objs, start=1):
        offsets.append(len(pdf))
        pdf += f"{index} 0 obj\n{obj}\nendobj\n".encode()
    xref_pos = len(pdf)
    pdf += f"xref\n0 {len(objs) + 1}\n0000000000 65535 f \n".encode()
    for offset in offsets:
        pdf += f"{offset:010d} 00000 n \n".encode()
    pdf += f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode()
    return pdf


def _docx_bytes() -> bytes:
    import io

    from docx import Document

    document = Document()
    document.add_heading("Q3 Research", level=1)
    document.add_paragraph(f"{DOCX_MARKER} {MARKER_ZH} 供应商分析正文。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Vendor"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = TABLE_MARKER
    table.cell(1, 1).text = "9.5"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


MARKDOWN_FIXTURE = f"""# 中文知识标题

{MARKER_ZH} 正文段落，包含唯一的中文检索标记。

## 数据表

| 指标 | 数值 |
| --- | --- |
| 唯一单元格 | {TABLE_MARKER} |
| 语言 | 中文 |
"""


@pytest.fixture
async def complete_schema(owner_engine):
    async with owner_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)


async def _seed_owner(owner_sessionmaker) -> tuple[uuid.UUID, uuid.UUID]:
    tenant_id = uuid.uuid4()
    owner_id = uuid.uuid4()
    suffix = uuid.uuid4().hex[:8]
    async with owner_sessionmaker() as session:
        session.add(Tenant(id=tenant_id, name="T", slug=f"pkbv-{suffix}"))
        await session.flush()
        session.add(
            User(
                id=owner_id,
                tenant_id=tenant_id,
                username=f"pkbv-{suffix}",
                email=f"pkbv-{suffix}@example.com",
                password_hash="not-a-real-password",
                display_name="PKB Vertical",
                role="member",
                is_active=True,
            )
        )
        await session.commit()
    return tenant_id, owner_id


class _session_context:
    def __init__(self, sessionmaker):
        self._sessionmaker = sessionmaker
        self._session = None

    async def __aenter__(self):
        self._session = self._sessionmaker()
        return self._session

    async def __aexit__(self, exc_type, exc, tb):
        if self._session is not None:
            await self._session.close()
        return False


async def _queue_and_process(
    owner_sessionmaker,
    service: PersonalKnowledgeService,
    *,
    tenant_id: uuid.UUID,
    owner_id: uuid.UUID,
) -> object:
    summary = await service.process_import_jobs(
        None,
        session_factory=lambda: _session_context(owner_sessionmaker),
        tenant_id=tenant_id,
        owner_user_id=owner_id,
        current_user_id=owner_id,
        limit=5,
        statuses=("queued",),
    )
    assert summary.failed == 0, summary.results
    assert summary.succeeded >= 1, summary.results
    return summary


async def test_pdf_docx_txt_markdown_import_vertical_evidence(complete_schema, owner_sessionmaker, tmp_path):
    tenant_id, owner_id = await _seed_owner(owner_sessionmaker)
    service = PersonalKnowledgeService(data_root=tmp_path)

    async with owner_sessionmaker() as session:
        pdf_result = await service.queue_source_bytes_import(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            filename="evidence.pdf",
            data=_minimal_pdf_bytes(
                [f"{PDF_MARKER} personal knowledge body", f"{MARKER_EN} shared english second line"]
            ),
            title="PDF evidence",
            source_kind="upload",
            source_uri="upload://evidence.pdf",
            created_by_user_id=owner_id,
            agent_searchable=True,
            sensitivity="internal",
            source_mime_type="application/pdf",
        )
        docx_result = await service.queue_source_bytes_import(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            filename="evidence.docx",
            data=_docx_bytes(),
            title="DOCX evidence",
            source_kind="upload",
            source_uri="upload://evidence.docx",
            created_by_user_id=owner_id,
            agent_searchable=True,
            sensitivity="internal",
            source_mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        txt_result = await service.queue_source_bytes_import(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            filename="evidence.txt",
            data=f"Plain text evidence.\n\n{TXT_MARKER} plain text body with {TXT_CELL} cell.\n".encode(),
            title="TXT evidence",
            source_kind="upload",
            source_uri="upload://evidence.txt",
            created_by_user_id=owner_id,
            agent_searchable=True,
            sensitivity="internal",
            source_mime_type="text/plain",
        )
        md_result = await service.queue_markdown_import(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            title="Markdown evidence",
            markdown=MARKDOWN_FIXTURE,
            source_kind="paste",
            source_uri=None,
            created_by_user_id=owner_id,
            agent_searchable=True,
            sensitivity="internal",
        )
        await session.commit()

    assert pdf_result.status == "queued"
    assert docx_result.status == "queued"
    assert txt_result.status == "queued"
    assert md_result.status == "queued"
    await _queue_and_process(owner_sessionmaker, service, tenant_id=tenant_id, owner_id=owner_id)

    async with owner_sessionmaker() as session:
        documents = (
            (
                await session.execute(
                    select(KnowledgeDocument)
                    .where(KnowledgeDocument.tenant_id == tenant_id)
                    .order_by(KnowledgeDocument.created_at.asc())
                )
            )
            .scalars()
            .all()
        )
        assert len(documents) == 4
        per_document_content: dict[str, str] = {}
        segment_snapshots: list[tuple[str | None, list, str, str]] = []
        for document in documents:
            assert document.status in {"ready", "degraded"}, document.status
            segments = (
                (
                    await session.execute(
                        select(KnowledgeSegment)
                        .where(KnowledgeSegment.document_id == document.id)
                        .order_by(KnowledgeSegment.position.asc())
                    )
                )
                .scalars()
                .all()
            )
            assert len(segments) > 0, document.title
            positions = [segment.position for segment in segments]
            assert positions == sorted(positions)
            for segment in segments:
                assert segment.segment_hash
                assert segment.document_id == document.id
                # Snapshot inside the session: detached instances cannot be
                # refreshed after close.
                segment_snapshots.append(
                    (str(document.id), list(segment.heading_path_json or []), segment.content, segment.position)
                )
            per_document_content[str(document.id)] = "\n".join(segment.content for segment in segments)
        await session.rollback()

    # Per-format proof: each document's own segments carry its distinct
    # marker (and table content where the fixture has one) — no format can
    # mask another format's loss through one aggregate string.
    assert PDF_MARKER in per_document_content[str(pdf_result.document_id)]
    assert MARKER_EN in per_document_content[str(pdf_result.document_id)]
    assert DOCX_MARKER in per_document_content[str(docx_result.document_id)]
    assert TABLE_MARKER in per_document_content[str(docx_result.document_id)]
    assert TXT_MARKER in per_document_content[str(txt_result.document_id)]
    assert TXT_CELL in per_document_content[str(txt_result.document_id)]
    assert MARKER_ZH in per_document_content[str(md_result.document_id)]
    assert TABLE_MARKER in per_document_content[str(md_result.document_id)]
    # The Markdown fixture's heading structure is preserved in heading_path.
    md_segments_with_heading = [snapshot for snapshot in segment_snapshots if snapshot[1] and MARKER_ZH in snapshot[2]]
    assert md_segments_with_heading, "expected the zh marker segment to carry its heading path"

    # Shared browser search returns the markers with document/segment citations.
    async with owner_sessionmaker() as session:
        en_hits = await service.search_personal(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            query=MARKER_EN,
            principal=HumanBrowserPrincipal(user_id=owner_id),
            limit=5,
        )
        zh_hits = await service.search_personal(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            query=MARKER_ZH,
            principal=HumanBrowserPrincipal(user_id=owner_id),
            limit=5,
        )
        await session.rollback()
    assert en_hits and all(MARKER_EN in hit.snippet or MARKER_EN in hit.title for hit in en_hits)
    assert all(hit.document_id and hit.segment_id for hit in en_hits)
    assert zh_hits and all(MARKER_ZH in hit.snippet or MARKER_ZH in hit.title for hit in zh_hits)


async def test_archived_document_leaves_search_results(complete_schema, owner_sessionmaker, tmp_path):
    tenant_id, owner_id = await _seed_owner(owner_sessionmaker)
    service = PersonalKnowledgeService(data_root=tmp_path)

    async with owner_sessionmaker() as session:
        result = await service.queue_markdown_import(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            title="Archive evidence",
            markdown=f"# Archive\n\n{MARKER_EN} archive lifecycle body",
            source_kind="paste",
            source_uri=None,
            created_by_user_id=owner_id,
            agent_searchable=True,
            sensitivity="internal",
        )
        await session.commit()
    await _queue_and_process(owner_sessionmaker, service, tenant_id=tenant_id, owner_id=owner_id)

    async with owner_sessionmaker() as session:
        hits = await service.search_personal(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            query=MARKER_EN,
            principal=HumanBrowserPrincipal(user_id=owner_id),
            limit=5,
        )
        await session.rollback()
    assert hits

    async with owner_sessionmaker() as session:
        archived = await service.patch_personal_document(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            document_id=result.document_id,
            current_user_id=owner_id,
            agent_id=None,
            status="archived",
        )
        await session.commit()
    assert archived is not None and archived.status == "archived"

    async with owner_sessionmaker() as session:
        hits_after = await service.search_personal(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            query=MARKER_EN,
            principal=HumanBrowserPrincipal(user_id=owner_id),
            limit=5,
        )
        await session.rollback()
    assert hits_after == []


async def test_corrupt_pdf_is_stable_typed_failure_not_raw_exception_prose(
    complete_schema, owner_sessionmaker, tmp_path
):
    """A corrupt PDF must land as a stable typed failure code on the job row —
    no raw exception prose as the user-facing error state."""
    from app.models.knowledge import KnowledgeIndexJob

    tenant_id, owner_id = await _seed_owner(owner_sessionmaker)
    service = PersonalKnowledgeService(data_root=tmp_path)

    async with owner_sessionmaker() as session:
        result = await service.queue_source_bytes_import(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            filename="corrupt.pdf",
            data=b"%PDF-1.4 this is not a real pdf body",
            title="Corrupt PDF",
            source_kind="upload",
            source_uri="upload://corrupt.pdf",
            created_by_user_id=owner_id,
            agent_searchable=True,
            sensitivity="internal",
            source_mime_type="application/pdf",
        )
        await session.commit()
    await _queue_and_process_expect_failure(owner_sessionmaker, service, tenant_id=tenant_id, owner_id=owner_id)

    async with owner_sessionmaker() as session:
        job = (
            await session.execute(select(KnowledgeIndexJob).where(KnowledgeIndexJob.id == result.job_id))
        ).scalar_one()
        snapshot = (
            str(job.status),
            str(job.error_message or ""),
            dict(job.job_metadata_json or {}),
        )
        await session.rollback()

    status, error_message, metadata = snapshot
    assert status == "failed"
    assert metadata.get("error") == "conversion_failed"
    # The user-facing error state is the typed code, not exception prose.
    assert error_message == "conversion_failed"


async def _queue_and_process_expect_failure(
    owner_sessionmaker, service, *, tenant_id: uuid.UUID, owner_id: uuid.UUID
) -> None:
    summary = await service.process_import_jobs(
        None,
        session_factory=lambda: _session_context(owner_sessionmaker),
        tenant_id=tenant_id,
        owner_user_id=owner_id,
        current_user_id=owner_id,
        limit=1,
        statuses=("queued",),
    )
    assert summary.succeeded == 0, summary.results


async def test_canonical_source_file_markers_source_ref_and_agent_runtime_citations(
    complete_schema, owner_sessionmaker, tmp_path
):
    """B9 vertical evidence: the canonical stored Markdown exists on disk with
    the expected marker/table content; the exact source_ref is preserved from
    segment store to search hit; a real DB AgentRuntimePrincipal can search
    and read with citations — no mocks at the acceptance boundary."""
    from app.models.agent import Agent
    from app.services.personal_knowledge_access import AgentRuntimePrincipal

    tenant_id, owner_id = await _seed_owner(owner_sessionmaker)
    agent_id = uuid.uuid4()
    async with owner_sessionmaker() as session:
        session.add(Agent(id=agent_id, name="PKB Vertical Agent", creator_id=owner_id, tenant_id=tenant_id))
        await session.commit()

    service = PersonalKnowledgeService(data_root=tmp_path)
    async with owner_sessionmaker() as session:
        md_result = await service.queue_markdown_import(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            title="Agent consumption evidence",
            markdown=MARKDOWN_FIXTURE,
            source_kind="paste",
            source_uri=None,
            created_by_user_id=owner_id,
            agent_searchable=True,
            sensitivity="internal",
        )
        await session.commit()
    await _queue_and_process(owner_sessionmaker, service, tenant_id=tenant_id, owner_id=owner_id)

    # 1. The canonical stored source file exists and carries marker + table.
    async with owner_sessionmaker() as session:
        document = (
            await session.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == md_result.document_id))
        ).scalar_one()
        canonical_rel = str(document.canonical_md_path or "")
        assert canonical_rel
        canonical_path = service.data_root / canonical_rel
        assert canonical_path.exists(), f"canonical Markdown missing on disk: {canonical_path}"
        canonical_text = canonical_path.read_text(encoding="utf-8")
        assert MARKER_ZH in canonical_text
        assert TABLE_MARKER in canonical_text
        # Segment rows carry stable identity for citation cross-checks.
        segments = (
            (
                await session.execute(
                    select(KnowledgeSegment)
                    .where(KnowledgeSegment.document_id == document.id)
                    .order_by(KnowledgeSegment.position.asc())
                )
            )
            .scalars()
            .all()
        )
        assert segments
        segment_ids = {str(segment.id) for segment in segments}
        await session.rollback()

    # 2. A real AgentRuntimePrincipal searches with authority and receives
    # citations whose source_ref matches the canonical contract exactly.
    principal = AgentRuntimePrincipal(
        agent_id=agent_id,
        requester_user_id=owner_id,
        session_id=str(uuid.uuid4()),
        purpose="interactive_session",
    )
    async with owner_sessionmaker() as session:
        search_result = await service.search_personal_with_authority(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            query=MARKER_ZH,
            principal=principal,
            limit=5,
        )
        await session.rollback()
    assert search_result.status == "ok", search_result.authority.deny_reason_code
    assert search_result.authority.allowed is True
    assert search_result.hits, "agent runtime principal must find the zh marker"
    for hit in search_result.hits:
        expected_ref = f"kb://person/{owner_id}/documents/{hit.document_id}#segment={hit.segment_id}"
        assert hit.source_ref == expected_ref
        assert str(hit.document_id) == str(md_result.document_id)
        assert str(hit.segment_id) in segment_ids
        assert hit.heading_path

    # 3. The same principal reads the document with citations intact.
    async with owner_sessionmaker() as session:
        read_result = await service.get_personal_document_with_authority(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            document_id=md_result.document_id,
            principal=principal,
        )
        await session.rollback()
    assert read_result.status == "ok", read_result.authority.deny_reason_code
    assert read_result.document is not None
    read_text = "\n".join(segment.content for segment in read_result.document.segments)
    assert MARKER_ZH in read_text
    assert TABLE_MARKER in read_text

    # 4. An agent owned by someone else gets a typed denial, not content.
    other_agent_id = uuid.uuid4()
    other_user_id = uuid.uuid4()
    async with owner_sessionmaker() as session:
        suffix = uuid.uuid4().hex[:8]
        session.add(
            User(
                id=other_user_id,
                tenant_id=tenant_id,
                username=f"pkbv-other-{suffix}",
                email=f"pkbv-other-{suffix}@example.com",
                password_hash="not-a-real-password",
                display_name="Other",
                role="member",
                is_active=True,
            )
        )
        await session.flush()
        session.add(Agent(id=other_agent_id, name="Other Agent", creator_id=other_user_id, tenant_id=tenant_id))
        await session.commit()
    denied_principal = AgentRuntimePrincipal(
        agent_id=other_agent_id,
        requester_user_id=other_user_id,
        session_id=str(uuid.uuid4()),
        purpose="interactive_session",
    )
    async with owner_sessionmaker() as session:
        denied = await service.search_personal_with_authority(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            query=MARKER_ZH,
            principal=denied_principal,
            limit=5,
        )
        await session.rollback()
    assert denied.authority.allowed is False
    assert denied.hits == []
    assert denied.authority.deny_reason_code


async def test_archived_document_authority_boundary_agent_vs_browser(complete_schema, owner_sessionmaker, tmp_path):
    """Archive is a real consumption boundary for agents, not for the owner workbench."""
    from app.models.agent import Agent
    from app.services.personal_knowledge_access import AgentRuntimePrincipal

    tenant_id, owner_id = await _seed_owner(owner_sessionmaker)
    agent_id = uuid.uuid4()
    async with owner_sessionmaker() as session:
        session.add(Agent(id=agent_id, name="Archive Boundary Agent", creator_id=owner_id, tenant_id=tenant_id))
        await session.commit()

    service = PersonalKnowledgeService(data_root=tmp_path)
    async with owner_sessionmaker() as session:
        result = await service.queue_markdown_import(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            title="Authority boundary doc",
            markdown=f"# Boundary\n\n{MARKER_EN} authority boundary body.",
            source_kind="paste",
            source_uri=None,
            created_by_user_id=owner_id,
            agent_searchable=True,
            sensitivity="internal",
        )
        await session.commit()
    await _queue_and_process(owner_sessionmaker, service, tenant_id=tenant_id, owner_id=owner_id)
    document_id = result.document_id
    agent_principal = AgentRuntimePrincipal(
        agent_id=agent_id,
        requester_user_id=owner_id,
        session_id=str(uuid.uuid4()),
        purpose="interactive_session",
    )
    browser_principal = HumanBrowserPrincipal(user_id=owner_id)

    # Baseline before archive: the agent reads with citations.
    async with owner_sessionmaker() as session:
        baseline_read = await service.get_personal_document_with_authority(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            document_id=document_id,
            principal=agent_principal,
        )
        await session.rollback()
    assert baseline_read.status == "ok"
    assert baseline_read.document is not None and baseline_read.document.segments

    async with owner_sessionmaker() as session:
        archived = await service.patch_personal_document(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            document_id=document_id,
            current_user_id=owner_id,
            agent_id=None,
            status="archived",
        )
        await session.commit()
    assert archived is not None and archived.status == "archived"

    # After archive: browser and agent search are both empty.
    async with owner_sessionmaker() as session:
        browser_hits = await service.search_personal(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            query=MARKER_EN,
            principal=browser_principal,
            limit=5,
        )
        agent_search = await service.search_personal_with_authority(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            query=MARKER_EN,
            principal=agent_principal,
            limit=5,
        )
        await session.rollback()
    assert browser_hits == []
    assert agent_search.hits == []

    # The agent read path returns no document/content for the archived
    # document; the owner browser keeps full read access for Restore.
    async with owner_sessionmaker() as session:
        agent_read = await service.get_personal_document_with_authority(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            document_id=document_id,
            principal=agent_principal,
        )
        browser_read = await service.get_personal_document_with_authority(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            document_id=document_id,
            principal=browser_principal,
        )
        await session.rollback()
    assert agent_read.document is None
    assert agent_read.status in {"empty", "denied", "unavailable"}
    assert browser_read.status == "ok"
    assert browser_read.document is not None
    assert MARKER_EN in "\n".join(segment.content for segment in browser_read.document.segments)

    # Restore brings the consumable state back; the agent reads with the
    # exact source_ref citation again.
    async with owner_sessionmaker() as session:
        restored = await service.restore_personal_document(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            document_id=document_id,
            current_user_id=owner_id,
        )
        await session.commit()
    assert restored is not None and restored.status in {"ready", "degraded"}
    async with owner_sessionmaker() as session:
        agent_read_after = await service.get_personal_document_with_authority(
            session,
            tenant_id=tenant_id,
            owner_user_id=owner_id,
            document_id=document_id,
            principal=agent_principal,
        )
        await session.rollback()
    assert agent_read_after.status == "ok"
    assert agent_read_after.document is not None
    assert agent_read_after.document.source_ref == f"kb://person/{owner_id}/documents/{document_id}"
    assert agent_read_after.document.segments
