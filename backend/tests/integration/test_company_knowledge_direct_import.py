"""Real-PostgreSQL vertical evidence for the Company Knowledge direct file import.

RC-02: admin multipart upload → durable job → real DocumentConversionService
(PDF/DOCX/TXT/Markdown with distinct per-format markers) → canonical evidence
chain → pre-publication segment preview → explicit idempotent proposal →
review/publish → namespace grant → member/Agent consumption with citations →
revoke/retire/restore boundaries. Lifecycle: committed running visibility,
final attempt, crash-at-cap terminalization, two-worker claim fencing,
cancel/retry, corrupt-file typed failure.
"""

from __future__ import annotations

import asyncio
import json
import uuid
from datetime import datetime, timedelta, timezone

import pytest
from sqlalchemy import func, select

from app.database import Base
from app.models.company_knowledge import (
    CompanyKnowledgeImportJob,
)
from app.models.knowledge import KnowledgeDocument, KnowledgeSegment
from app.models.tenant import Tenant
from app.models.user import User
from app.services.company_knowledge_contracts import SourceContractInput
from app.services.company_knowledge_control_plane import (
    CompanyKnowledgePermissionGrantInput,
    CompanyKnowledgePermissionService,
)
from app.services.company_knowledge_gateway import (
    CompanyKnowledgeGateway,
    CompanyKnowledgeReadRequest,
    CompanyKnowledgeSearchRequest,
    CompanyKnowledgeSourceExplainRequest,
)
from app.services.company_knowledge_permissions import CompanyKnowledgePrincipal
from app.services.company_knowledge_service import (
    CompanyKnowledgeImportError,
    CompanyKnowledgeJobConflict,
    CompanyKnowledgeReviewRequest,
    CompanyKnowledgeService,
)

MARKER_PDF = "WEEKEND-RC-20260825-CKB-PDF-MARKER"
MARKER_DOCX = "WEEKEND-RC-20260825-CKB-DOCX-MARKER"
MARKER_TXT = "WEEKEND-RC-20260825-CKB-TXT-MARKER"
MARKER_ZH = "周末RC20260825公司知识唯一标记"
TABLE_MARKER = "CKB-TABLE-CELL-MARKER-6613"
NAMESPACE = "company/general"


def _minimal_pdf_bytes(lines: list[str]) -> bytes:
    """Hand-crafted minimal PDF (standard Helvetica; ASCII text only)."""
    content = [f"BT /F1 12 Tf 72 {720 - 20 * index} Td ({line}) Tj ET" for index, line in enumerate(lines)]
    objs = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
    ]
    stream = "\n".join(content)
    objs.append(f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream")
    objs.append("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    pdf = b"%PDF-1.4\n"
    offsets: list[int] = []
    for index, obj in enumerate(objs, start=1):
        offsets.append(len(pdf))
        pdf += f"{index} 0 obj\n{obj}\nendobj\n".encode()
    xref_pos = len(pdf)
    pdf += f"xref\n0 {len(objs) + 1}\n0000000000 65535 f \n".encode()
    for offset in offsets:
        pdf += f"{offset:010d} 00000 n \n".encode()
    pdf += f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode()
    return pdf


def _docx_bytes() -> bytes:
    import io

    from docx import Document

    document = Document()
    document.add_heading("Company Runbook", level=1)
    document.add_paragraph(f"{MARKER_DOCX} {MARKER_ZH} 公司运维手册正文。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "System"
    table.cell(0, 1).text = "Owner"
    table.cell(1, 0).text = TABLE_MARKER
    table.cell(1, 1).text = "oncall"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


MARKDOWN_FIXTURE = f"""# 公司知识标题

{MARKER_ZH} 正文段落，包含唯一的中文检索标记。

## 数据表

| 指标 | 数值 |
| --- | --- |
| 唯一单元格 | {TABLE_MARKER} |
| 语言 | 中文 |
"""


@pytest.fixture
async def complete_schema(owner_engine):
    async with owner_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)


def _principal(*, tenant_id: uuid.UUID, user_id: uuid.UUID) -> CompanyKnowledgePrincipal:
    return CompanyKnowledgePrincipal(
        tenant_id=tenant_id,
        accountable_user_id=user_id,
        accountable_role="org_admin",
        actor_type="user",
        actor_id=user_id,
        purpose="interactive_session",
        session_id="company-kb-direct-import",
    )


def _member_principal(*, tenant_id: uuid.UUID, user_id: uuid.UUID) -> CompanyKnowledgePrincipal:
    return CompanyKnowledgePrincipal(
        tenant_id=tenant_id,
        accountable_user_id=user_id,
        accountable_role="member",
        actor_type="user",
        actor_id=user_id,
        purpose="interactive_session",
        session_id="company-kb-member",
    )


def _managed_file_contract() -> SourceContractInput:
    return SourceContractInput(
        source_kind="managed_file",
        provider_kind="manual_upload",
        stable_source_id="company-file-upload",
        owner_principal_ref="role:org_admin",
        accountable_steward_ref="role:org_admin",
        connection_ref=None,
        schema_ref=None,
        schema_version=None,
        identity_keys=("source_item_id",),
        relation_keys=(),
        ingest_mode="manual",
        cursor_kind=None,
        cursor_policy={},
        watermark_field=None,
        temporal_mapping={"observed_at": "ingest_time"},
        source_acl_mapping_policy={"mode": "required_snapshot"},
        default_sensitivity="PL1_public",
        export_policy={"allowed": False},
        retention_policy={"class": "company_record"},
        legal_hold_policy={"supported": True},
        allowed_namespaces=(NAMESPACE,),
        precedence_policy_ref=None,
        acceptance_suite_ref=None,
        idempotency_policy={"key": "source_item_id+revision"},
    )


async def _seed_tenant(owner_sessionmaker) -> dict[str, uuid.UUID]:
    ids = {key: uuid.uuid4() for key in ("tenant", "admin", "member", "agent")}
    suffix = uuid.uuid4().hex[:8]
    async with owner_sessionmaker() as session:
        session.add(Tenant(id=ids["tenant"], name="Company", slug=f"ckb-{suffix}"))
        await session.flush()
        session.add(
            User(
                id=ids["admin"],
                tenant_id=ids["tenant"],
                username=f"ckb-admin-{suffix}",
                email=f"ckb-admin-{suffix}@example.com",
                password_hash="not-a-real-password",
                display_name="CKB Admin",
                role="org_admin",
                is_active=True,
            )
        )
        session.add(
            User(
                id=ids["member"],
                tenant_id=ids["tenant"],
                username=f"ckb-member-{suffix}",
                email=f"ckb-member-{suffix}@example.com",
                password_hash="not-a-real-password",
                display_name="CKB Member",
                role="member",
                is_active=True,
            )
        )
        await session.flush()
        from app.models.agent import Agent

        session.add(Agent(id=ids["agent"], name="CKB Agent", creator_id=ids["admin"], tenant_id=ids["tenant"]))
        from app.models.security_audit import ResourcePermission

        session.add(
            ResourcePermission(
                tenant_id=ids["tenant"],
                principal_type="user",
                principal_id=ids["admin"],
                resource_type="company_knowledge_scope",
                resource_id=ids["tenant"],
                actions=[
                    "approve",
                    "publish",
                    "retire",
                    "restore",
                    "discover",
                    "search",
                    "read",
                    "cite",
                    "propose",
                    "review",
                ],
                conditions={},
                effect="allow",
                sensitivity_ceiling="PL3_sensitive",
                purposes=["interactive_session"],
                created_by_user_id=ids["admin"],
            )
        )
        await session.commit()
    return ids


async def _register_contract(owner_sessionmaker, service, *, tenant_id, admin_id):
    async with owner_sessionmaker() as session:
        contract = await service.register_source_contract(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            contract_input=_managed_file_contract(),
            idempotency_key=f"contract-{uuid.uuid4().hex[:12]}",
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
        return contract


async def _queue_file(
    owner_sessionmaker,
    service,
    *,
    tenant_id,
    admin_id,
    contract,
    filename,
    data,
    title,
    mime,
    idempotency_key,
):
    async with owner_sessionmaker() as session:
        job = await service.queue_direct_file_import(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            source_contract_id=contract.id,
            source_contract_version=contract.version,
            filename=filename,
            data=data,
            source_mime_type=mime,
            title=title,
            proposed_namespace=NAMESPACE,
            proposed_sensitivity="internal",
            purpose="RC-02 vertical evidence",
            source_acl_snapshot={"all_tenant_members": True},
            idempotency_key=idempotency_key,
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
        return job


async def _process(owner_sessionmaker, service, *, tenant_id, job_id):
    return await service.process_import_job(
        tenant_id=tenant_id,
        job_id=job_id,
        session_factory=owner_sessionmaker,
    )


async def _read_job(owner_sessionmaker, job_id: uuid.UUID) -> tuple[str, int, str | None]:
    async with owner_sessionmaker() as session:
        job = (
            await session.execute(select(CompanyKnowledgeImportJob).where(CompanyKnowledgeImportJob.id == job_id))
        ).scalar_one()
        snapshot = (str(job.status or ""), int(job.attempt_count or 0), job.last_error_code)
        await session.rollback()
        return snapshot


async def test_pdf_docx_txt_markdown_direct_import_vertical_evidence(complete_schema, owner_sessionmaker, tmp_path):
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id, member_id, agent_id = ids["tenant"], ids["admin"], ids["member"], ids["agent"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)

    jobs = {}
    fixtures = {
        "pdf": (
            "evidence.pdf",
            _minimal_pdf_bytes([f"{MARKER_PDF} company body", "second pdf line"]),
            "application/pdf",
        ),
        "docx": (
            "evidence.docx",
            _docx_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "txt": ("evidence.txt", f"Plain company note.\n\n{MARKER_TXT} plain text body.\n".encode(), "text/plain"),
        "md": ("evidence.md", MARKDOWN_FIXTURE.encode("utf-8"), "text/markdown"),
    }
    for key, (filename, data, mime) in fixtures.items():
        jobs[key] = await _queue_file(
            owner_sessionmaker,
            service,
            tenant_id=tenant_id,
            admin_id=admin_id,
            contract=contract,
            filename=filename,
            data=data,
            title=f"{key.upper()} evidence",
            mime=mime,
            idempotency_key=f"ckb-vertical-{key}-{uuid.uuid4().hex[:8]}",
        )
        assert jobs[key].status == "queued"

    for key, job in jobs.items():
        processed = await _process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=job.id)
        assert processed.status == "completed", (key, processed.last_error_code, processed.last_error)
        assert processed.document_id is not None and processed.evidence_id is not None
        jobs[key] = processed

    # Per-format proof: each document's own segments carry its distinct marker.
    async with owner_sessionmaker() as session:
        for key, marker in (("pdf", MARKER_PDF), ("docx", MARKER_DOCX), ("txt", MARKER_TXT), ("md", MARKER_ZH)):
            document = (
                await session.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == jobs[key].document_id))
            ).scalar_one()
            assert document.status == "ready", (key, document.status)
            canonical_path = service._data_root / str(document.canonical_md_path)
            assert canonical_path.exists(), f"canonical Markdown missing: {canonical_path}"
            canonical_text = canonical_path.read_text(encoding="utf-8")
            assert marker in canonical_text, key
            segments = (
                (
                    await session.execute(
                        select(KnowledgeSegment)
                        .where(KnowledgeSegment.document_id == document.id)
                        .order_by(KnowledgeSegment.position.asc())
                    )
                )
                .scalars()
                .all()
            )
            assert len(segments) > 0, key
            positions = [segment.position for segment in segments]
            assert positions == sorted(positions)
            joined = "\n".join(segment.content for segment in segments)
            assert marker in joined, key
            if key in {"docx", "md"}:
                assert TABLE_MARKER in joined, key
        await session.rollback()

    # Pre-publication preview: admin-only segment view with provenance refs.
    md_job = jobs["md"]
    async with owner_sessionmaker() as session:
        preview = await service.get_import_job_preview(
            session,
            tenant_id=tenant_id,
            job_id=md_job.id,
        )
        await session.rollback()
    assert preview is not None
    assert preview.document_id == md_job.document_id
    assert preview.evidence_id == md_job.evidence_id
    assert preview.source_id == md_job.source_id
    assert preview.segments, "preview must expose converted segments"
    assert any(MARKER_ZH in segment.content for segment in preview.segments)
    assert all(segment.heading_path is not None and segment.token_count >= 0 for segment in preview.segments)

    # Explicit proposal: created once, submitted, origin=direct_import; the
    # second call returns the same proposal (idempotent, no duplicate).
    async with owner_sessionmaker() as session:
        proposal = await service.create_proposal_from_import(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            job_id=md_job.id,
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
    assert proposal.status == "submitted"
    assert proposal.proposal_kind == "knowledge"
    assert proposal.source_document_id == md_job.document_id
    assert str(dict(proposal.proposed_patch_json or {}).get("origin") or "") == "direct_import"

    async with owner_sessionmaker() as session:
        proposal_again = await service.create_proposal_from_import(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            job_id=md_job.id,
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
    assert proposal_again.id == proposal.id

    # Review → publish through the existing authority (normal risk: one
    # org_admin approval is enough).
    async with owner_sessionmaker() as session:
        reviewed = await service.record_review(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            proposal_id=proposal.id,
            request=CompanyKnowledgeReviewRequest(
                decision="approve",
                reviewer_role="org_admin",
                reason="RC-02 vertical approval",
                evidence_refs=(f"company-evidence://{proposal.source_refs_json[0].split('://')[-1]}",),
                policy_snapshot={"policy": "rc02-vertical-v1"},
            ),
            expected_state_version=proposal.state_version,
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
    assert reviewed.status == "approved"

    async with owner_sessionmaker() as session:
        publication = await service.publish_proposal(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            proposal_id=proposal.id,
            expected_state_version=reviewed.state_version,
            valid_from=datetime.now(timezone.utc),
            valid_until=None,
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
    assert publication.status == "active"

    # Namespace grants to the member and the agent; both consume with citations.
    permission_service = CompanyKnowledgePermissionService(proposal_authority=service)
    member_grant_request = CompanyKnowledgePermissionGrantInput(
        principal_type="user",
        principal_id=member_id,
        principal_key=None,
        resource_type="company_knowledge_namespace",
        resource_id=None,
        resource_key=f"namespace:{NAMESPACE}",
        actions=("discover", "search", "read", "cite"),
        effect="allow",
        sensitivity_ceiling="PL3_sensitive",
        purposes=("interactive_session",),
        expires_at=None,
        idempotency_key=f"grant-member-{uuid.uuid4().hex[:8]}",
    )
    agent_grant_request = CompanyKnowledgePermissionGrantInput(
        principal_type="agent",
        principal_id=agent_id,
        principal_key=None,
        resource_type="company_knowledge_namespace",
        resource_id=None,
        resource_key=f"namespace:{NAMESPACE}",
        actions=("discover", "search", "read", "cite"),
        effect="allow",
        sensitivity_ceiling="PL3_sensitive",
        purposes=("interactive_session",),
        expires_at=None,
        idempotency_key=f"grant-agent-{uuid.uuid4().hex[:8]}",
    )
    async with owner_sessionmaker() as session:
        member_grant = await permission_service.grant_permission(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            request=member_grant_request,
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await permission_service.grant_permission(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            request=agent_grant_request,
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()

    gateway = CompanyKnowledgeGateway()
    async with owner_sessionmaker() as session:
        member_search = await gateway.search(
            session,
            principal=_member_principal(tenant_id=tenant_id, user_id=member_id),
            request=CompanyKnowledgeSearchRequest(
                query=MARKER_ZH,
                filters={"namespaces": [NAMESPACE]},
                limit=5,
                trace_id=f"trace-{uuid.uuid4().hex[:12]}",
            ),
        )
        await session.rollback()
    assert member_search.results, "member must find the zh marker after grant"
    assert all(hit.document_id and hit.segment_id for hit in member_search.results)
    assert any(str(publication.id) == str(hit.publication_id) for hit in member_search.results)
    assert all(hit.source_ref for hit in member_search.results)

    agent_principal = CompanyKnowledgePrincipal(
        tenant_id=tenant_id,
        accountable_user_id=admin_id,
        accountable_role="org_admin",
        actor_type="agent",
        actor_id=agent_id,
        purpose="interactive_session",
        session_id="company-kb-agent-session",
    )
    async with owner_sessionmaker() as session:
        agent_read = await gateway.read(
            session,
            principal=agent_principal,
            request=CompanyKnowledgeReadRequest(
                document_id=publication.document_id,
                publication_id=publication.id,
                segment_ids=(),
                max_chars=4000,
                trace_id=f"trace-{uuid.uuid4().hex[:12]}",
            ),
        )
        await session.rollback()
    assert agent_read.status == "ok", agent_read.status
    assert agent_read.citations, "agent read must carry citations"
    assert MARKER_ZH in "\n".join(segment.content for segment in agent_read.segments)

    # Explain: the direct-import evidence explains with source_ref, coverage,
    # and the import receipt — never the canonical path or source bytes.
    async with owner_sessionmaker() as session:
        explained = await gateway.explain_source(
            session,
            principal=agent_principal,
            request=CompanyKnowledgeSourceExplainRequest(
                evidence_id=md_job.evidence_id,
                trace_id=f"trace-{uuid.uuid4().hex[:12]}",
            ),
        )
        await session.rollback()
    assert explained.status == "ok", explained.status
    payload = dict(explained.payload or {})
    assert payload.get("source_ref") == f"company-evidence://{md_job.evidence_id}"
    assert payload.get("publication_id") == str(publication.id)
    assert dict(payload.get("coverage") or {}).get("complete") is True
    assert payload.get("ingestion_receipt_ref") == f"company-import://{md_job.id}"
    serialized = json.dumps(payload, ensure_ascii=False)
    assert "canonical_md_path" not in serialized
    assert str(tmp_path) not in serialized
    assert MARKER_ZH not in serialized

    # Revoke → member loses discovery immediately; retire → agent loses read;
    # restore → consumption returns under the still-live grant.
    async with owner_sessionmaker() as session:
        await permission_service.revoke_permission(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            permission_id=uuid.UUID(str(member_grant["permission_id"])),
            reason="RC-02 vertical revoke",
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
    async with owner_sessionmaker() as session:
        member_search_after = await gateway.search(
            session,
            principal=_member_principal(tenant_id=tenant_id, user_id=member_id),
            request=CompanyKnowledgeSearchRequest(
                query=MARKER_ZH,
                filters={"namespaces": [NAMESPACE]},
                limit=5,
                trace_id=f"trace-{uuid.uuid4().hex[:12]}",
            ),
        )
        await session.rollback()
    assert member_search_after.results == ()

    async with owner_sessionmaker() as session:
        await service.retire_publication(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            publication_id=publication.id,
            reason="RC-02 vertical retire",
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
    async with owner_sessionmaker() as session:
        agent_read_after_retire = await gateway.read(
            session,
            principal=agent_principal,
            request=CompanyKnowledgeReadRequest(
                document_id=publication.document_id,
                publication_id=publication.id,
                segment_ids=(),
                max_chars=4000,
                trace_id=f"trace-{uuid.uuid4().hex[:12]}",
            ),
        )
        await session.rollback()
    assert agent_read_after_retire.status != "ok"

    async with owner_sessionmaker() as session:
        restored = await service.restore_publication(
            session,
            principal=_principal(tenant_id=tenant_id, user_id=admin_id),
            publication_id=publication.id,
            reason="RC-02 vertical restore",
            valid_from=datetime.now(timezone.utc),
            trace_id=f"trace-{uuid.uuid4().hex[:12]}",
        )
        await session.commit()
    assert restored.status == "active"
    async with owner_sessionmaker() as session:
        agent_read_restored = await gateway.read(
            session,
            principal=agent_principal,
            request=CompanyKnowledgeReadRequest(
                document_id=restored.document_id,
                publication_id=restored.id,
                segment_ids=(),
                max_chars=4000,
                trace_id=f"trace-{uuid.uuid4().hex[:12]}",
            ),
        )
        await session.rollback()
    assert agent_read_restored.status == "ok"
    assert MARKER_ZH in "\n".join(segment.content for segment in agent_read_restored.segments)


async def test_direct_import_claim_commits_running_before_conversion(complete_schema, owner_sessionmaker, tmp_path):
    """A concurrent reader sees committed running + attempt 1 while conversion is still in flight."""
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id = ids["tenant"], ids["admin"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)
    job = await _queue_file(
        owner_sessionmaker,
        service,
        tenant_id=tenant_id,
        admin_id=admin_id,
        contract=contract,
        filename="evidence.md",
        data=f"# Visible\n\n{MARKER_ZH} visible claim body.".encode(),
        title="Visible claim",
        mime="text/markdown",
        idempotency_key=f"ckb-visible-{uuid.uuid4().hex[:8]}",
    )

    entered = asyncio.Event()
    release = asyncio.Event()
    loop = asyncio.get_running_loop()

    class _GatedConverter:
        def convert_bytes(self, **kwargs):
            loop.call_soon_threadsafe(entered.set)
            import time

            deadline = time.monotonic() + 30
            while not release.is_set() and time.monotonic() < deadline:
                time.sleep(0.01)
            from app.services.document_conversion import DocumentConversionResult

            return DocumentConversionResult(
                markdown="# Visible\n\nconverted",
                plain_text="Visible converted",
                source_path="",
                source_uri=None,
                source_sha256="c" * 64,
                source_mime_type="text/markdown",
                engine="gated-test",
                used_ocr=False,
                used_vision=False,
                page_count=1,
                artifact_markdown_path="",
                artifact_metadata_path="",
                warnings=(),
            )

    service._conversion_service_override = _GatedConverter()
    worker = asyncio.create_task(_process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=job.id))
    try:
        await asyncio.wait_for(entered.wait(), timeout=30)
        status, attempts, _ = await _read_job(owner_sessionmaker, job.id)
        assert status == "running", status
        assert attempts == 1, attempts
    finally:
        release.set()
        processed = await asyncio.wait_for(worker, timeout=60)
    assert processed.status == "completed"


async def test_final_attempt_executes_and_crash_at_cap_terminalizes(complete_schema, owner_sessionmaker, tmp_path):
    """attempt==max-1 runs its final claim; a crash at attempt==max with an expired claim is terminalized, not rerun."""
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id = ids["tenant"], ids["admin"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)

    final_job = await _queue_file(
        owner_sessionmaker,
        service,
        tenant_id=tenant_id,
        admin_id=admin_id,
        contract=contract,
        filename="final.md",
        data=f"# Final\n\n{MARKER_ZH} final attempt body.".encode(),
        title="Final attempt",
        mime="text/markdown",
        idempotency_key=f"ckb-final-{uuid.uuid4().hex[:8]}",
    )
    async with owner_sessionmaker() as session:
        job_row = (
            await session.execute(select(CompanyKnowledgeImportJob).where(CompanyKnowledgeImportJob.id == final_job.id))
        ).scalar_one()
        job_row.attempt_count = final_job.max_attempts - 1
        await session.commit()
    processed = await _process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=final_job.id)
    assert processed.status == "completed"
    assert processed.attempt_count == final_job.max_attempts

    crashed_job = await _queue_file(
        owner_sessionmaker,
        service,
        tenant_id=tenant_id,
        admin_id=admin_id,
        contract=contract,
        filename="crashed.md",
        data=f"# Crash\n\n{MARKER_ZH} crash body.".encode(),
        title="Crashed worker",
        mime="text/markdown",
        idempotency_key=f"ckb-crash-{uuid.uuid4().hex[:8]}",
    )
    async with owner_sessionmaker() as session:
        job_row = (
            await session.execute(
                select(CompanyKnowledgeImportJob).where(CompanyKnowledgeImportJob.id == crashed_job.id)
            )
        ).scalar_one()
        job_row.status = "running"
        job_row.attempt_count = crashed_job.max_attempts
        job_row.claim_token = uuid.uuid4()
        job_row.claim_expires_at = datetime.now(timezone.utc) - timedelta(minutes=10)
        await session.commit()

    conversion_calls: list[str] = []

    class _CountingConverter:
        def convert_bytes(self, **kwargs):  # pragma: no cover - must not run
            conversion_calls.append("convert")
            raise AssertionError("exhausted stale claim must not rerun conversion")

    service._conversion_service_override = _CountingConverter()
    async with owner_sessionmaker() as session:
        summary = await service.recover_due_import_jobs(
            session,
            session_factory=owner_sessionmaker,
            limit=10,
        )
        await session.commit()
    assert conversion_calls == []
    assert summary.attempted >= 1
    assert summary.failed >= 1
    status, attempts, error_code = await _read_job(owner_sessionmaker, crashed_job.id)
    assert status == "failed", status
    assert attempts == crashed_job.max_attempts
    assert error_code == "company_knowledge_import_attempts_exhausted", error_code


async def test_two_workers_cannot_process_same_job_twice(complete_schema, owner_sessionmaker, tmp_path):
    """Concurrent claims of one job produce exactly one winner; the loser gets a typed claim outcome and no duplicate document exists."""
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id = ids["tenant"], ids["admin"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)
    job = await _queue_file(
        owner_sessionmaker,
        service,
        tenant_id=tenant_id,
        admin_id=admin_id,
        contract=contract,
        filename="race.md",
        data=f"# Race\n\n{MARKER_ZH} race body.".encode(),
        title="Race",
        mime="text/markdown",
        idempotency_key=f"ckb-race-{uuid.uuid4().hex[:8]}",
    )

    outcomes: list[tuple[str, str]] = []

    async def run_worker(name: str) -> None:
        try:
            await _process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=job.id)
            outcomes.append((name, "completed"))
        except RuntimeError as exc:
            outcomes.append((name, str(exc)))

    await asyncio.gather(run_worker("a"), run_worker("b"), run_worker("a"))
    completed = [name for name, outcome in outcomes if outcome == "completed"]
    assert len(completed) == 1, outcomes
    typed_losses = {
        "company_knowledge_import_job_already_claimed",
        "company_knowledge_import_claim_lost",
    }
    assert all(outcome in typed_losses or outcome == "completed" for _name, outcome in outcomes), outcomes

    async with owner_sessionmaker() as session:
        documents = (
            await session.execute(
                select(func.count())
                .select_from(KnowledgeDocument)
                .where(
                    KnowledgeDocument.tenant_id == tenant_id,
                    KnowledgeDocument.scope_type == "company",
                )
            )
        ).scalar_one()
        await session.rollback()
    assert int(documents) == 1


async def test_failed_job_terminalizes_at_ceiling_and_is_not_reselected(complete_schema, owner_sessionmaker, tmp_path):
    """A deterministically failing job exhausts bounded attempts into terminal failed and is never silently reselected."""
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id = ids["tenant"], ids["admin"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)
    job = await _queue_file(
        owner_sessionmaker,
        service,
        tenant_id=tenant_id,
        admin_id=admin_id,
        contract=contract,
        filename="always-corrupt.pdf",
        data=b"%PDF-1.4 not a real pdf body",
        title="Always corrupt",
        mime="application/pdf",
        idempotency_key=f"ckb-corrupt-{uuid.uuid4().hex[:8]}",
    )
    for _ in range(job.max_attempts):
        with pytest.raises(CompanyKnowledgeImportError) as exc:
            await _process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=job.id)
        assert exc.value.code == "conversion_failed", exc.value.code
        status, attempts, _ = await _read_job(owner_sessionmaker, job.id)
        if status == "failed":
            break
        # Bring the backoff forward so the next drain/claim is due now.
        async with owner_sessionmaker() as session:
            job_row = (
                await session.execute(select(CompanyKnowledgeImportJob).where(CompanyKnowledgeImportJob.id == job.id))
            ).scalar_one()
            job_row.available_at = datetime.now(timezone.utc) - timedelta(seconds=1)
            await session.commit()

    status, attempts, error_code = await _read_job(owner_sessionmaker, job.id)
    assert status == "failed", (status, attempts)
    assert attempts == job.max_attempts
    assert error_code == "conversion_failed", error_code

    async with owner_sessionmaker() as session:
        summary = await service.recover_due_import_jobs(
            session,
            session_factory=owner_sessionmaker,
            limit=10,
        )
        await session.commit()
    assert summary.attempted == 0, summary.job_refs


async def test_cancel_and_retry_lifecycle(complete_schema, owner_sessionmaker, tmp_path):
    """Cancel is a queued-only CAS with a committed cancelled_at; retry requeues failed/cancelled under the ceiling and rejects permanent/at-cap jobs."""
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id = ids["tenant"], ids["admin"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)
    job = await _queue_file(
        owner_sessionmaker,
        service,
        tenant_id=tenant_id,
        admin_id=admin_id,
        contract=contract,
        filename="cancel.md",
        data=f"# Cancel\n\n{MARKER_ZH} cancel body.".encode(),
        title="Cancel",
        mime="text/markdown",
        idempotency_key=f"ckb-cancel-{uuid.uuid4().hex[:8]}",
    )

    async with owner_sessionmaker() as session:
        cancelled = await service.cancel_import_job(
            session,
            tenant_id=tenant_id,
            job_id=job.id,
        )
        await session.commit()
    assert cancelled.lifecycle_status == "cancelled"
    assert cancelled.cancelled_at

    async with owner_sessionmaker() as session:
        with pytest.raises(CompanyKnowledgeJobConflict) as conflict:
            await service.cancel_import_job(session, tenant_id=tenant_id, job_id=job.id)
        await session.rollback()
    assert conflict.value.code == "not_cancellable_terminal"

    async with owner_sessionmaker() as session:
        retried = await service.retry_import_job(
            session,
            tenant_id=tenant_id,
            job_id=job.id,
        )
        await session.commit()
    assert retried.lifecycle_status == "queued"
    assert retried.error_code is None
    assert retried.cancelled_at is None

    # A terminal failed job at the attempt ceiling refuses requeue.
    async with owner_sessionmaker() as session:
        job_row = (
            await session.execute(select(CompanyKnowledgeImportJob).where(CompanyKnowledgeImportJob.id == job.id))
        ).scalar_one()
        job_row.status = "failed"
        job_row.attempt_count = job_row.max_attempts
        await session.commit()
    async with owner_sessionmaker() as session:
        with pytest.raises(CompanyKnowledgeJobConflict) as cap_conflict:
            await service.retry_import_job(session, tenant_id=tenant_id, job_id=job.id)
        await session.rollback()
    assert cap_conflict.value.code == "retry_attempt_limit"


async def test_corrupt_pdf_typed_failure_and_reupload_recovers(complete_schema, owner_sessionmaker, tmp_path):
    """A corrupt PDF lands as typed conversion_failed (retryable); re-uploading clean bytes under a new key completes and becomes searchable."""
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id = ids["tenant"], ids["admin"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)

    bad_job = await _queue_file(
        owner_sessionmaker,
        service,
        tenant_id=tenant_id,
        admin_id=admin_id,
        contract=contract,
        filename="corrupt.pdf",
        data=b"%PDF-1.4 this is not a real pdf body",
        title="Corrupt PDF",
        mime="application/pdf",
        idempotency_key=f"ckb-corrupt2-{uuid.uuid4().hex[:8]}",
    )
    with pytest.raises(CompanyKnowledgeImportError) as first_failure:
        await _process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=bad_job.id)
    assert first_failure.value.code == "conversion_failed", first_failure.value.code
    # After the first failure the job sits in bounded-backoff queued with the
    # typed code visible (the drain owns automatic retries); terminal failure
    # at the ceiling is not manually retryable.
    status, attempts_after_first, error_code = await _read_job(owner_sessionmaker, bad_job.id)
    assert status == "queued", status
    assert attempts_after_first < bad_job.max_attempts
    assert error_code == "conversion_failed", error_code

    async with owner_sessionmaker() as session:
        summary = await service.get_import_job_summary(
            session,
            tenant_id=tenant_id,
            job_id=bad_job.id,
        )
        await session.rollback()
    assert summary is not None
    assert summary.error_code == "conversion_failed"
    for _ in range(bad_job.max_attempts):
        status, _attempts, _ = await _read_job(owner_sessionmaker, bad_job.id)
        if status == "failed":
            break
        async with owner_sessionmaker() as session:
            job_row = (
                await session.execute(
                    select(CompanyKnowledgeImportJob).where(CompanyKnowledgeImportJob.id == bad_job.id)
                )
            ).scalar_one()
            job_row.available_at = datetime.now(timezone.utc) - timedelta(seconds=1)
            await session.commit()
        with pytest.raises(CompanyKnowledgeImportError) as loop_failure:
            await _process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=bad_job.id)
        assert loop_failure.value.code == "conversion_failed", loop_failure.value.code
    status, attempts, error_code = await _read_job(owner_sessionmaker, bad_job.id)
    assert status == "failed", (status, attempts)
    assert attempts == bad_job.max_attempts
    assert error_code == "conversion_failed"
    async with owner_sessionmaker() as session:
        terminal_summary = await service.get_import_job_summary(
            session,
            tenant_id=tenant_id,
            job_id=bad_job.id,
        )
        await session.rollback()
    assert terminal_summary is not None
    assert terminal_summary.terminal is True
    assert terminal_summary.retryable is False

    good_job = await _queue_file(
        owner_sessionmaker,
        service,
        tenant_id=tenant_id,
        admin_id=admin_id,
        contract=contract,
        filename="recovered.pdf",
        data=_minimal_pdf_bytes([f"{MARKER_PDF} recovered clean pdf body"]),
        title="Recovered PDF",
        mime="application/pdf",
        idempotency_key=f"ckb-recovered-{uuid.uuid4().hex[:8]}",
    )
    processed = await _process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=good_job.id)
    assert processed.status == "completed", processed.last_error

    async with owner_sessionmaker() as session:
        # Pre-publication content is not publicly searchable; prove the
        # segments exist via the preview read model instead.
        preview = await service.get_import_job_preview(
            session,
            tenant_id=tenant_id,
            job_id=good_job.id,
        )
        await session.rollback()
    assert preview is not None
    assert any(MARKER_PDF in segment.content for segment in preview.segments)


async def test_non_admin_principal_cannot_queue_direct_import(complete_schema, owner_sessionmaker, tmp_path):
    """A plain member principal is denied at the service permission gate."""
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id, member_id = ids["tenant"], ids["admin"], ids["member"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)

    async with owner_sessionmaker() as session:
        with pytest.raises(PermissionError):
            await service.queue_direct_file_import(
                session,
                principal=_member_principal(tenant_id=tenant_id, user_id=member_id),
                source_contract_id=contract.id,
                source_contract_version=contract.version,
                filename="denied.md",
                data=b"# denied",
                source_mime_type="text/markdown",
                title="Denied",
                proposed_namespace=NAMESPACE,
                proposed_sensitivity="internal",
                purpose="denied",
                source_acl_snapshot={"all_tenant_members": True},
                idempotency_key=f"ckb-denied-{uuid.uuid4().hex[:8]}",
                trace_id=f"trace-{uuid.uuid4().hex[:12]}",
            )
        await session.rollback()


async def test_direct_import_idempotency_hash_covers_semantic_inputs(complete_schema, owner_sessionmaker, tmp_path):
    """Same key with changed title/purpose/MIME conflicts; an identical replay returns the same job."""
    ids = await _seed_tenant(owner_sessionmaker)
    tenant_id, admin_id = ids["tenant"], ids["admin"]
    service = CompanyKnowledgeService(data_root=tmp_path)
    contract = await _register_contract(owner_sessionmaker, service, tenant_id=tenant_id, admin_id=admin_id)
    principal = _principal(tenant_id=tenant_id, user_id=admin_id)
    key = f"ckb-idem-{uuid.uuid4().hex[:8]}"
    data = f"# Idem\n\n{MARKER_ZH} idempotency body.".encode()

    async def queue_once(*, title, purpose, mime, idem_key=key):
        async with owner_sessionmaker() as session:
            job = await service.queue_direct_file_import(
                session,
                principal=principal,
                source_contract_id=contract.id,
                source_contract_version=contract.version,
                filename="idem.md",
                data=data,
                source_mime_type=mime,
                title=title,
                proposed_namespace=NAMESPACE,
                proposed_sensitivity="internal",
                purpose=purpose,
                source_acl_snapshot={"all_tenant_members": True},
                idempotency_key=idem_key,
                trace_id=f"trace-{uuid.uuid4().hex[:12]}",
            )
            await session.commit()
            return job

    first = await queue_once(title="Idem", purpose="original", mime="text/markdown")
    replay = await queue_once(title="Idem", purpose="original", mime="text/markdown")
    assert replay.id == first.id

    for changed in (
        {"title": "Idem changed"},
        {"purpose": "changed purpose"},
        {"mime": "text/plain"},
    ):
        kwargs = {"title": "Idem", "purpose": "original", "mime": "text/markdown", **changed}
        with pytest.raises(ValueError, match="company_knowledge_import_idempotency_conflict"):
            await queue_once(**kwargs)

    # Leave no recoverable queued row behind for the shared fleet-wide drain.
    processed = await _process(owner_sessionmaker, service, tenant_id=tenant_id, job_id=first.id)
    assert processed.status == "completed"
